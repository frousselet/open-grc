"""ISO 27005 risk assessment DOCX export.

Produces a structured Word document covering the full assessment cycle:
context, criteria, threat/vulnerability catalogs, ISO 27005 analyses,
consolidated risks, treatment plans and acceptances. DOCX helpers and
color palette are reused from `reports/management_review.py`.
"""

from io import BytesIO

from django.utils import timezone

from accounts.models import CompanySettings
from reports.management_review import (
    _CLR_ACCENT,
    _CLR_HEADING,
    _CLR_MUTED,
    _docx_add_kv_table,
    _docx_add_table,
    _docx_setup_styles,
)


def generate_iso27005_report_docx(assessment, user):
    """Generate an ISO 27005 risk assessment DOCX report.

    Sections:
      1. Context (assessment metadata, scopes)
      2. Risk criteria (likelihood/impact scales, risk levels, matrix)
      3. Threats referenced in the assessment
      4. Vulnerabilities referenced in the assessment
      5. ISO 27005 analyses (threat x vulnerability scenarios)
      6. Consolidated risks
      7. Treatment plans
      8. Risk acceptances

    Returns (filename, content_bytes).
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    from risks.models import RiskAcceptance, RiskTreatmentPlan, Threat, Vulnerability

    doc = Document()
    _docx_setup_styles(doc)

    now = timezone.now()
    try:
        company_name = CompanySettings.get().name or "Cairn"
    except Exception:
        company_name = "Cairn"

    # ── Cover page ──
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_name)
    run.font.size = Pt(28)
    run.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(*_CLR_HEADING)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p_fmt = p._p.get_or_add_pPr()
    p_bdr = p_fmt.makeelement(qn("w:pBdr"), {})
    bottom = p_bdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): f"{_CLR_ACCENT[0]:02X}{_CLR_ACCENT[1]:02X}{_CLR_ACCENT[2]:02X}",
    })
    p_bdr.append(bottom)
    p_fmt.append(p_bdr)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("ISO 27005 Risk Assessment Report")
    run.font.size = Pt(20)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(*_CLR_HEADING)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{assessment.reference} - {assessment.name}")
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(*_CLR_MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    generated_by = ""
    if user:
        generated_by = (
            getattr(user, "display_name", "") or getattr(user, "email", "") or ""
        )
    suffix = f" - {generated_by}" if generated_by else ""
    run = p.add_run(f"Generated on {now.strftime('%Y-%m-%d %H:%M')}{suffix}")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(*_CLR_MUTED)

    doc.add_page_break()

    # ── 1. Context ──
    doc.add_heading("1. Context", level=1)
    scope_names = ", ".join(s.name for s in assessment.scopes.all()) or "-"
    description = (assessment.description or "").strip()
    if len(description) > 300:
        description = description[:297] + "..."
    pairs = [
        ("Reference", str(assessment.reference or "-")),
        ("Name", str(assessment.name or "-")),
        ("Methodology", str(assessment.get_methodology_display())),
        ("Assessor", str(assessment.assessor) if assessment.assessor_id else "-"),
        (
            "Assessment date",
            assessment.assessment_date.isoformat()
            if assessment.assessment_date
            else "-",
        ),
        ("Status", str(assessment.get_status_display())),
        ("Scopes", scope_names),
        ("Description", description or "-"),
    ]
    _docx_add_kv_table(doc, pairs)

    # ── 2. Risk criteria ──
    doc.add_heading("2. Risk criteria", level=1)
    criteria = assessment.risk_criteria
    if criteria:
        doc.add_heading(f"{criteria.reference} - {criteria.name}", level=2)
        if criteria.acceptance_threshold:
            doc.add_paragraph(
                f"Acceptance threshold: {criteria.acceptance_threshold}",
            )

        doc.add_heading("Likelihood scale", level=3)
        l_rows = [
            [str(sl.level), sl.name, sl.description or ""]
            for sl in criteria.scale_levels.filter(scale_type="likelihood").order_by("level")
        ]
        _docx_add_table(doc, ["Level", "Name", "Description"], l_rows)

        doc.add_heading("Impact scale", level=3)
        i_rows = [
            [str(sl.level), sl.name, sl.description or ""]
            for sl in criteria.scale_levels.filter(scale_type="impact").order_by("level")
        ]
        _docx_add_table(doc, ["Level", "Name", "Description"], i_rows)

        doc.add_heading("Risk levels", level=3)
        r_rows = [
            [
                str(rl.level),
                rl.name,
                rl.description or "",
                "Yes" if rl.requires_treatment else "No",
            ]
            for rl in criteria.risk_levels.order_by("level")
        ]
        _docx_add_table(
            doc,
            ["Level", "Name", "Description", "Requires treatment"],
            r_rows,
        )

        doc.add_heading("Risk matrix", level=3)
        matrix = criteria.risk_matrix or {}
        if matrix:
            ls = sorted({int(k.split(",")[0]) for k in matrix.keys()})
            is_ = sorted({int(k.split(",")[1]) for k in matrix.keys()})
            headers = ["L \\ I"] + [str(i) for i in is_]
            rows = [
                [str(l)] + [str(matrix.get(f"{l},{i}", "-")) for i in is_]
                for l in ls
            ]
            _docx_add_table(doc, headers, rows)
        else:
            doc.add_paragraph("Empty matrix.")
    else:
        doc.add_paragraph("No risk criteria configured for this assessment.")

    # ── 3. Threats ──
    doc.add_heading("3. Threats", level=1)
    threat_ids = list(
        assessment.iso27005_risks.values_list("threat_id", flat=True).distinct()
    )
    threats = Threat.objects.filter(pk__in=threat_ids).order_by("reference")
    if threats:
        rows = [
            [
                t.reference,
                t.name,
                t.get_type_display() or "",
                t.get_origin_display() or "",
                str(t.typical_likelihood) if t.typical_likelihood is not None else "-",
            ]
            for t in threats
        ]
        _docx_add_table(
            doc,
            ["Ref", "Name", "Type", "Origin", "Typical likelihood"],
            rows,
        )
    else:
        doc.add_paragraph("No threats referenced in this assessment.")

    # ── 4. Vulnerabilities ──
    doc.add_heading("4. Vulnerabilities", level=1)
    vuln_ids = list(
        assessment.iso27005_risks.values_list("vulnerability_id", flat=True).distinct()
    )
    vulnerabilities = Vulnerability.objects.filter(pk__in=vuln_ids).order_by("reference")
    if vulnerabilities:
        rows = [
            [
                v.reference,
                v.name,
                v.get_category_display() or "",
                str(v.severity) if v.severity is not None else "-",
            ]
            for v in vulnerabilities
        ]
        _docx_add_table(
            doc,
            ["Ref", "Name", "Category", "Severity"],
            rows,
        )
    else:
        doc.add_paragraph("No vulnerabilities referenced in this assessment.")

    # ── 5. ISO 27005 analyses ──
    doc.add_heading("5. ISO 27005 analyses", level=1)
    analyses = (
        assessment.iso27005_risks.select_related("threat", "vulnerability", "risk")
        .order_by("reference")
    )
    if analyses.exists():
        rows = []
        for a in analyses:
            rows.append([
                a.reference,
                a.threat.name if a.threat_id else "-",
                a.vulnerability.name if a.vulnerability_id else "-",
                str(a.combined_likelihood) if a.combined_likelihood is not None else "-",
                str(a.max_impact) if a.max_impact is not None else "-",
                str(a.risk_level) if a.risk_level is not None else "-",
                a.risk.reference if a.risk_id else "-",
            ])
        _docx_add_table(
            doc,
            ["Ref", "Threat", "Vulnerability", "Likelihood", "Impact", "Level", "Consolidated risk"],
            rows,
        )
    else:
        doc.add_paragraph("No ISO 27005 analyses recorded.")

    # ── 6. Consolidated risks ──
    doc.add_heading("6. Consolidated risks", level=1)
    risks = assessment.risks.select_related("risk_owner").order_by("reference")
    if risks.exists():
        rows = []
        for r in risks:
            rows.append([
                r.reference,
                r.name,
                str(r.initial_risk_level) if r.initial_risk_level is not None else "-",
                str(r.current_risk_level) if r.current_risk_level is not None else "-",
                str(r.residual_risk_level) if r.residual_risk_level is not None else "-",
                r.get_treatment_decision_display() or "-",
                r.get_status_display(),
            ])
        _docx_add_table(
            doc,
            ["Ref", "Name", "Initial", "Current", "Residual", "Decision", "Status"],
            rows,
        )
    else:
        doc.add_paragraph("No consolidated risks.")

    # ── 7. Treatment plans ──
    doc.add_heading("7. Treatment plans", level=1)
    plans = (
        RiskTreatmentPlan.objects.filter(risk__assessment=assessment)
        .select_related("risk", "owner")
        .order_by("reference")
    )
    if plans.exists():
        rows = []
        for p in plans:
            rows.append([
                p.reference,
                p.risk.reference if p.risk_id else "-",
                p.name,
                p.get_treatment_type_display() or "-",
                p.get_status_display(),
                p.target_date.isoformat() if p.target_date else "-",
                f"{p.progress_percentage}%",
            ])
        _docx_add_table(
            doc,
            ["Ref", "Risk", "Name", "Type", "Status", "Target", "Progress"],
            rows,
        )
    else:
        doc.add_paragraph("No treatment plans.")

    # ── 8. Risk acceptances ──
    doc.add_heading("8. Risk acceptances", level=1)
    acceptances = (
        RiskAcceptance.objects.filter(risk__assessment=assessment)
        .select_related("risk", "accepted_by")
        .order_by("reference")
    )
    if acceptances.exists():
        rows = []
        for a in acceptances:
            accepted_by = (
                getattr(a.accepted_by, "display_name", None)
                or getattr(a.accepted_by, "email", None)
                or "-"
            )
            rows.append([
                a.reference,
                a.risk.reference if a.risk_id else "-",
                a.get_status_display(),
                str(a.risk_level_at_acceptance)
                if a.risk_level_at_acceptance is not None
                else "-",
                a.valid_until.isoformat() if a.valid_until else "-",
                accepted_by,
            ])
        _docx_add_table(
            doc,
            ["Ref", "Risk", "Status", "Level", "Valid until", "Accepted by"],
            rows,
        )
    else:
        doc.add_paragraph("No risk acceptances.")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    date_str = now.strftime("%Y%m%d_%H%M%S")
    filename = f"ISO27005_Report_{assessment.reference}_{date_str}.docx"
    return filename, buf.getvalue()
